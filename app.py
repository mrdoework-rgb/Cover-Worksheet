import io
import json
import os
import re
import tempfile
import urllib.error
import urllib.request
from pathlib import Path
from uuid import uuid4

import streamlit as st
import streamlit.components.v1 as components
from docx import Document
from docx.shared import Pt


DEFAULT_WORKSHEET_STRUCTURE = """Every worksheet must structurally contain the following content, mapped to the JSON schema in Part 2:
  1. Core Text: A very short paragraph or bulleted list explaining the absolute core facts/concepts. Only essential keywords should serve as technical vocabulary. If a core formula or relationship exists, display it prominently (e.g., $$Formula$$).
  2. Quick Questions: Exactly 8 binary-choice questions formatted strictly as "Does [Factor A] affect [Option X] or [Option Y]?" or "Is [Concept] an example of [Option X] or [Option Y]?". Answers must be directly extractable from the core text. Put an answer line after every quick question, using a short blank line or "____" marker so students can write their answers.
  3. True or False: Exactly 8 simple statement verification questions.
  4. Factor Summary Table: A comparison table with between 2 to 5 columns where students can classify 4 distinct items/factors from the text. Make about 60% of the cells empty so students can complete them, while keeping the table structure clear.
  5. Fill in the Blanks: A short paragraph cloze activity (3-4 sentences) with a clean Word Bank listed directly under it.
  6. Application Problems: Exactly 6 simple contextual questions that push students to apply the content in straightforward ways.
  7. Extension: Exactly 3 higher tier type contextual application questions pushing students to apply knowledge to novel scenarios, evaluate safety/experimental designs, or explain "why". Leave an empty line between the application questions and the extension questions.
"""

PROMPT_TEMPLATE = """You are an expert GCSE Science resource creator specializing in evidence-based pedagogies for lower-attaining and SEND students, combined with a precise JSON parsing architecture.

Your task is to generate a minimalist, highly structured cover lesson worksheet based on the topic provided below, and output it directly as a raw JSON object following a strict schema.

### TARGET TOPIC
Topic: {topic}
Year Group: {year_group}
Ability of the group: {ability}

### PART 1: PEDAGOGICAL & CONTENT CONSTRAINTS

Always adhere to the following layout, tone, and structural constraints when designing the content:
- Keep text to an absolute minimum. Use minimal instructions. Avoid dense walls of text. Use clean formatting to maximize scannability.
- Adapt the vocabulary, sentence length, and complexity to the specified year group and ability level. If the class is younger or lower-attaining, use simpler wording, shorter sentences, and more direct instructions. If the class is older or higher-attaining, you may use slightly more precise language.
- Make the core text, every question, and every instruction suitable for the class age and ability level.
- For the Application Problems and Extension sections, use the JSON type "long_question" for each question so they are visually separated with extra spacing in the Word document.
- The final part of the JSON must be an answer key. It should be a numbered list of all the answers and should use the JSON type "answer_key".
{worksheet_structure}

### PART 2: JSON SCHEMA & TECHNICAL CONSTRAINTS

1. NO MARKDOWN WRAPPERS: Output raw JSON only. Do not wrap the output in ```json blocks. Do not include any conversational text before or after the JSON object.
2. NO CUSTOM TYPES: You must ONLY use the types explicitly listed in the schema below. Do not use "list", "callout", "equation", "binary_choice", or "cloze". Map all text structures to these seven approved strings: "subheading", "paragraph", "table", "question", "long_question", "answer_key", or "image".
3. CHEMICAL FORMULA RULES: You must format all chemical formulas and numerical subscripts using HTML <sub> tags.
   - Standard Molecular: H<sub>2</sub>O, CO<sub>2</sub>
   - Polyatomic Ions/Parentheses: Ca(NO<sub>3</sub>)<sub>2</sub>
   - States of Matter: Write as standard text in parentheses, e.g., 2H<sub>2</sub>O(l)

### JSON Structure Definition:
- `worksheet_title`: (String) The main title of the resource.
- `document_blocks`: (Array of Objects) The sequential content of the document. EVERY object must have a `"type"` key. The `"type"` MUST be exactly one of these seven strings:
    - If type is "subheading", "paragraph", "question", "long_question", or "answer_key": Include a `"content"` key with the text. Apply Subscript Rules.
    - If type is "image": Include a `"content"` key containing a brief text description of a useful diagram for this topic.
    - If type is "table": Do NOT include a `"content"` key. Instead, include a `"table_data"` object containing:
        - `has_headers`: (Boolean)
        - `column_count`: (Integer)
        - `rows`: (Array of Arrays of Strings) Apply Subscript Rules to cell text.

### MAPPING GUIDE (How to construct the JSON from the parts):
- Use "subheading" blocks to transition between the mandatory parts (e.g., "content": "Part 1: Quick Questions").
- For "True or False" statements, map each statement to a "question" type, appending " [True / False]" to the end of the content string.
- For "Fill in the Blanks", output the paragraph with blanks written as "______" as a "paragraph" type, and output the Word Bank as a subsequent "paragraph" type.
- For all other questions, map them directly to individual "question" type objects.
- For the Application Problems and Extension sections, map each question to a "long_question" type object so they are separated with extra spacing in the Word document.
- Put the answer key as the final block in the JSON, using the type "answer_key" and a numbered list of answers in the content string.

Immediately output the complete raw JSON object for the topic: {topic}. Do not include conversational filler before or after the JSON."""


# --- 1. CHEMISTRY REGEX FUNCTION ---
def append_science_text(paragraph, raw_text):
    parts = re.split(r'(<sub>.*?</sub>)', raw_text)
    for part in parts:
        if part.startswith('<sub>') and part.endswith('</sub>'):
            subscript_content = part.replace('<sub>', '').replace('</sub>', '')
            run = paragraph.add_run(subscript_content)
            run.font.subscript = True
        else:
            paragraph.add_run(part)


def clean_question_text(raw_text):
    if not raw_text:
        return ""

    cleaned = re.sub(r'^\s*(?:\d+[.)]\s*|[A-Za-z][.)]\s*)+', '', str(raw_text))
    return cleaned.strip()


# --- 2. GEMINI API HELPERS ---
def build_gemini_prompt(topic, year_group=None, ability=None, worksheet_structure=None):
    year_group_text = (year_group or "unknown").strip()
    ability_text = (ability or "unknown").strip()
    structure_text = (worksheet_structure or DEFAULT_WORKSHEET_STRUCTURE).strip()
    return PROMPT_TEMPLATE.format(
        topic=topic.strip(),
        year_group=year_group_text,
        ability=ability_text,
        worksheet_structure=structure_text,
    )


def get_gemini_api_key():
    api_key = os.getenv("GEMINI_API_KEY")
    if api_key:
        return api_key

    try:
        return st.secrets["GEMINI_API_KEY"]
    except Exception:
        return None


def resolve_gemini_api_key():
    st.markdown("Watch this video to learn how to get a Free Gemini API key")
    components.iframe(
        "https://www.youtube.com/embed/Uyn-P2nRvDA",
        height=158,
        width=280,
        scrolling=False,
    )
    st.markdown(
        "Create your Google Gemini API Key 👉 "
        "[https://aistudio.google.com/](https://aistudio.google.com/)"
    )
    return st.text_input(
        "Gemini API key",
        type="password",
        placeholder="Enter your Gemini API key",
        help="This value is used for the request when provided.",
    )


def clean_json_text(raw_text):
    text = raw_text.strip()
    if text.startswith("```"):
        lines = text.splitlines()
        if len(lines) >= 2:
            text = "\n".join(lines[1:-1])
        else:
            text = text.strip("`")
    return text.strip()


def parse_gemini_response(response_data):
    try:
        text = response_data["candidates"][0]["content"]["parts"][0]["text"]
    except (KeyError, IndexError, TypeError) as exc:
        raise ValueError("Gemini response did not contain a usable content block.") from exc

    cleaned_text = clean_json_text(text)
    return json.loads(cleaned_text)


def call_gemini_api(topic, year_group=None, ability=None, api_key=None, worksheet_structure=None):
    if not topic or not topic.strip():
        raise ValueError("Please enter a topic first.")

    resolved_key = api_key or get_gemini_api_key()
    if not resolved_key:
        raise ValueError("Enter a Gemini API key or set the GEMINI_API_KEY environment variable or Streamlit secret before generating a worksheet.")

    model = os.getenv("GEMINI_MODEL") or "gemini-2.5-flash"
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={resolved_key}"
    payload = {
        "contents": [
            {
                "role": "user",
                "parts": [{"text": build_gemini_prompt(topic, year_group=year_group, ability=ability, worksheet_structure=worksheet_structure)}],
            }
        ],
        "generationConfig": {"temperature": 0.1},
    }

    request = urllib.request.Request(
        url,
        data=json.dumps(payload).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST",
    )

    try:
        with urllib.request.urlopen(request, timeout=60) as response:
            response_data = json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="ignore")
        raise RuntimeError(f"Gemini API request failed ({exc.code}): {detail}") from exc
    except urllib.error.URLError as exc:
        raise RuntimeError(f"Unable to reach Gemini API: {exc}") from exc

    return parse_gemini_response(response_data)


# --- 3. V3 COMPONENT DOCUMENT COMPILER ---
def generate_worksheet(json_data, template_path):
    template = Path(template_path)
    doc = Document(template) if template.exists() else Document()
    title = json_data.get("worksheet_title", "Science Worksheet")
    blocks = json_data.get("document_blocks", [])

    for p in doc.paragraphs:
        # --- A. TITLE ---
        if "[[WORKSHEET_TITLE]]" in p.text:
            p.text = ""
            append_science_text(p, title)

        # --- B. THE DYNAMIC CONTENT BLOCK ---
        elif "[[CONTENT_HERE]]" in p.text:
            for block in blocks:
                b_type = block.get("type", "unknown")

                if b_type == "subheading":
                    new_p = p.insert_paragraph_before('')
                    new_p.style = 'Heading 2'
                    append_science_text(new_p, block.get("content", ""))

                elif b_type == "paragraph":
                    new_p = p.insert_paragraph_before('')
                    new_p.style = 'Normal'
                    append_science_text(new_p, block.get("content", ""))

                elif b_type == "question":
                    new_p = p.insert_paragraph_before('')
                    new_p.style = 'List Number'
                    append_science_text(new_p, clean_question_text(block.get("content", "")))

                elif b_type == "long_question":
                    new_p = p.insert_paragraph_before('')
                    new_p.style = 'List Number'
                    new_p.paragraph_format.space_after = Pt(24)
                    append_science_text(new_p, clean_question_text(block.get("content", "")))

                elif b_type == "answer_key":
                    continue

                elif b_type == "image":
                    new_p = p.insert_paragraph_before('')
                    new_p.style = 'Normal'
                    image_desc = block.get("content", "Insert image here")
                    run = new_p.add_run(f"[ 📷 PLACEHOLDER: {image_desc} ]")
                    run.font.bold = True

                elif b_type == "table":
                    table_data = block.get("table_data", {})
                    rows_data = table_data.get("rows", [])

                    if rows_data:
                        col_count = table_data.get("column_count", len(rows_data[0]))
                        row_count = len(rows_data)
                        table = doc.add_table(rows=row_count, cols=col_count)
                        try:
                            table.style = 'Science_Table_Style'
                        except KeyError:
                            pass

                        p._p.addprevious(table._tbl)

                        for r_idx, row_data in enumerate(rows_data):
                            for c_idx, cell_text in enumerate(row_data):
                                cell_paragraph = table.cell(r_idx, c_idx).paragraphs[0]
                                cell_paragraph.text = ""
                                append_science_text(cell_paragraph, str(cell_text))

                else:
                    new_p = p.insert_paragraph_before('')
                    new_p.style = 'Normal'
                    error_text = str(block.get("content", f"[AI Generated Unrecognized Block: '{b_type}']"))
                    append_science_text(new_p, error_text)

            p._element.getparent().remove(p._element)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer, title


def generate_answer_key_document(json_data, template_path):
    template = Path(template_path)
    doc = Document(template) if template.exists() else Document()
    title = json_data.get("worksheet_title", "Science Worksheet")
    blocks = json_data.get("document_blocks", [])
    answer_blocks = [block for block in blocks if block.get("type") == "answer_key"]

    for p in doc.paragraphs:
        if "[[WORKSHEET_TITLE]]" in p.text:
            p.text = ""
            append_science_text(p, f"{title} - Answers")

        elif "[[CONTENT_HERE]]" in p.text:
            for block in answer_blocks:
                new_p = p.insert_paragraph_before('')
                new_p.style = 'Normal'
                append_science_text(new_p, block.get("content", ""))

            p._element.getparent().remove(p._element)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer, f"{title} - Answers"


def prepare_template_path(uploaded_template):
    default_template = Path("Generic Template.docx")
    if uploaded_template is None:
        if default_template.exists():
            return str(default_template)
        raise FileNotFoundError("The default template file 'Generic Template.docx' was not found.")

    temp_dir = Path(tempfile.gettempdir()) / "cover_worksheet_templates"
    temp_dir.mkdir(exist_ok=True)
    safe_name = Path(uploaded_template.name).stem.replace(" ", "_")
    temp_path = temp_dir / f"{safe_name}_{uuid4().hex}.docx"
    temp_path.write_bytes(uploaded_template.getvalue())
    return str(temp_path)


# --- 4. STREAMLIT UI ---
def main():
    st.set_page_config(page_title="Science Worksheet Generator", page_icon="🔬")

    st.title("Worksheet Generator")
    st.write(
        "This app helps you create a ready-to-use worksheet in just a few steps. You type in the topic, year group, and ability level, "
        "and the app sends that information to Gemini, which writes the worksheet content in a structured format. Streamlit is the tool "
        "that runs the website and collects your inputs. The app then turns that content into a Word document you can download. "
        "GitHub stores the project files and keeps the app organised, so it can be shared and updated easily."
    )

    topic_input = st.text_input("What should this worksheet be about", placeholder="e.g. Photosynthesis")
    year_group_input = st.text_input("Year Group", placeholder="e.g. Year 10")
    ability_input = st.text_input("What is the ability of the group?", placeholder="e.g. Lower-attaining / mixed / high ability")
    worksheet_structure_input = st.text_area(
        "Worksheet elements",
        value=DEFAULT_WORKSHEET_STRUCTURE,
        height=220,
        help="Edit the worksheet structure here before generating the worksheet.",
    )
    api_key_input = resolve_gemini_api_key()
    uploaded_template = st.file_uploader("Optional: upload a Word template (.docx)", type=["docx"])

    if "worksheet_bytes" not in st.session_state:
        st.session_state["worksheet_bytes"] = None
    if "worksheet_file_name" not in st.session_state:
        st.session_state["worksheet_file_name"] = "worksheet.docx"
    if "answer_bytes" not in st.session_state:
        st.session_state["answer_bytes"] = None
    if "answer_file_name" not in st.session_state:
        st.session_state["answer_file_name"] = "answers.docx"

    with st.expander("Template guidance", expanded=False):
        st.write(
            "This app creates a cover lesson worksheet from your topic, year group, and ability level. It sends the request to Gemini, "
            "builds a Word document, and can also generate a separate answers document. The app is hosted in the cloud and the code is "
            "stored in GitHub."
        )
        st.table(
            {
                "Style": ["Heading 2", "Normal", "List Number"],
                "Used for": [
                    "Section headings such as Quick Questions and Extension",
                    "Paragraph text, instructions, and image placeholders",
                    "Question numbering and list-style questions",
                ],
            }
        )

        template_path = Path("Generic Template.docx")
        if template_path.exists():
            template_bytes = template_path.read_bytes()
            st.download_button(
                label="Download starter template",
                data=template_bytes,
                file_name="Generic Template.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    with st.expander("Environmental Impact", expanded=False):
        st.write(
            "Each use of this app consumes around 1000 input tokens and 5000 output tokens. "
            "That is roughly the same energy use as watching about half an hour of YouTube on a laptop, "
            "driving about 150m in an electric car, or boiling enough water to make one cup of tea."
        )

    if st.button("Generate Worksheet", type="primary"):
        if topic_input.strip():
            try:
                with st.spinner("Contacting Gemini to build the worksheet..."):
                    json_data = call_gemini_api(
                        topic_input,
                        year_group=year_group_input,
                        ability=ability_input,
                        api_key=api_key_input or None,
                        worksheet_structure=worksheet_structure_input,
                    )
                template_path = prepare_template_path(uploaded_template)
                word_buffer, doc_title = generate_worksheet(json_data, template_path)
                answer_buffer, answer_title = generate_answer_key_document(json_data, template_path)
                st.session_state["worksheet_bytes"] = word_buffer.getvalue()
                st.session_state["worksheet_file_name"] = f"{doc_title}.docx"
                st.session_state["answer_bytes"] = answer_buffer.getvalue()
                st.session_state["answer_file_name"] = f"{answer_title} - Answers.docx"
                st.success("Worksheet generated successfully!")
            except Exception as exc:
                st.error(f"❌ An error occurred: {exc}")
        else:
            st.error("Please enter a topic first.")

    if st.session_state.get("worksheet_bytes") is not None:
        st.download_button(
            label="📥 Download Word Document",
            data=st.session_state["worksheet_bytes"],
            file_name=st.session_state["worksheet_file_name"],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    if st.session_state.get("answer_bytes") is not None:
        st.download_button(
            label="📥 Download Answers Document",
            data=st.session_state["answer_bytes"],
            file_name=st.session_state["answer_file_name"],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


if __name__ == "__main__":
    main()
